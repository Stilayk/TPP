from __future__ import annotations

import io
import re
from datetime import date, datetime
from pathlib import Path
from urllib.parse import quote

from fastapi import HTTPException
from openpyxl import Workbook
from openpyxl.styles import Alignment

from app.config import settings
from app.employee_exit_blocks import compose_employee_exit_instruction
from app.models import DailyReport, User
from app.schemas import DailyReportOut, ReportEntryOut, UserOut

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[misc, assignment]

EXPORT_FILENAME_RE = re.compile(r"^report_(\d+)_\d{4}-\d{2}-\d{2}\.xlsx$")

_WS_INLINE = re.compile(r"[ \t]{2,}")


def normalize_employee_exit_instruction_text(text: str) -> str:
    """Убирает лишние пустые строки и схлопывает повторяющиеся пробелы в строках (общий вид и Word)."""
    text = (text or "").strip()
    if not text:
        return ""
    chunks: list[str] = []
    for raw_chunk in re.split(r"\n\s*\n+", text):
        lines: list[str] = []
        for raw_ln in raw_chunk.splitlines():
            ln = raw_ln.strip()
            if not ln:
                continue
            ln = _WS_INLINE.sub(" ", ln).strip()
            if ln:
                lines.append(ln)
        if lines:
            chunks.append("\n".join(lines))
    return "\n\n".join(chunks)


def is_safe_export_filename(filename: str) -> bool:
    if not filename or "/" in filename or "\\" in filename:
        return False
    return bool(re.fullmatch(r"[A-Za-z0-9._-]+", filename))


def exports_dir() -> Path:
    p = Path(settings.EXPORTS_DIR).expanduser()
    if not p.is_absolute():
        p = (Path.cwd() / p).resolve()
    p.mkdir(parents=True, exist_ok=True)
    return p


def build_employee_exit_instruction(
    fio: str, login: str, password: str, domain: str, blocks: list[str] | None = None
) -> str:
    raw = compose_employee_exit_instruction(fio, login, password, domain, blocks=blocks)
    return normalize_employee_exit_instruction_text(raw)


def build_employee_exit_instruction_docx_bytes(
    fio: str, login: str, password: str, domain: str, blocks: list[str] | None = None
) -> bytes:
    if Document is None:
        raise HTTPException(
            status_code=503,
            detail="Генерация Word недоступна: в образе не установлен пакет python-docx. Пересоберите backend.",
        )
    from docx.enum.text import WD_BREAK, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def _apply_instruction_run_font(run) -> None:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except (AttributeError, TypeError):
            pass

    text = build_employee_exit_instruction(fio, login, password, domain, blocks=blocks)
    doc = Document()
    for block in (b.strip() for b in text.split("\n\n")):
        if not block:
            continue
        lines = [_WS_INLINE.sub(" ", ln.strip()) for ln in block.split("\n") if ln.strip()]
        if not lines:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        first = True
        for ln in lines:
            if not first:
                br = p.add_run()
                br.add_break(WD_BREAK.LINE)
                _apply_instruction_run_font(br)
            r = p.add_run(ln)
            _apply_instruction_run_font(r)
            first = False
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_export_name_part(value: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "_", value).strip()
    return cleaned or "Сотрудник"


def attachment_filename_docx(base_name: str) -> str:
    safe = safe_export_name_part(base_name) or "employee"
    return f"instrukciya_{safe}.docx"


def attachment_content_disposition_docx(filename: str) -> str:
    ascii_fallback = re.sub(r"[^A-Za-z0-9._-]+", "_", filename) or "instrukciya.docx"
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{quote(filename)}"


def report_excel_filename(report: DailyReport) -> str:
    return f"report_{report.id}_{report.date.isoformat()}.xlsx"


def surname_from_full_name(full_name: str) -> str:
    parts = [p for p in (full_name or "").strip().split() if p]
    if not parts:
        return "Сотрудник"
    return parts[0]


def format_date_ru(d: date) -> str:
    return d.strftime("%d.%m.%Y")


def minutes_ru_label(minutes: int) -> str:
    n = int(minutes)
    if n == 0:
        return "0 минут"
    mod100 = n % 100
    if 11 <= mod100 <= 14:
        return f"{n} минут"
    last = n % 10
    if last == 1:
        return f"{n} минута"
    if last in (2, 3, 4):
        return f"{n} минуты"
    return f"{n} минут"


def build_report_excel(report: DailyReport, user: User) -> Path:
    out_dir = exports_dir()
    filename = report_excel_filename(report)
    out_path = out_dir / filename

    wb = Workbook()
    ws = wb.active
    ws.title = "Отчёт"

    ws.merge_cells("B1:C1")
    ws["A1"] = "дата"
    ws["B1"] = format_date_ru(report.date)
    ws["B1"].alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells("B2:C2")
    ws["A2"] = "ФИО"
    ws["B2"] = (user.full_name or "").strip() or user.username
    ws["B2"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ws["A3"] = "Задача"
    ws["B3"] = "Результат"
    ws["C3"] = "Время работы"

    row = 4
    for entry in report.entries:
        ws[f"A{row}"] = entry.task or ""
        ws[f"B{row}"] = entry.description
        ws[f"C{row}"] = minutes_ru_label(entry.minutes)
        row += 1

    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 18

    wb.save(out_path)
    return out_path


def report_to_out(db, report: DailyReport) -> DailyReportOut:
    employee = db.get(User, report.support_user_id)
    if not employee:
        raise HTTPException(status_code=500, detail="Employee not found")
    return DailyReportOut(
        report_id=report.id,
        date=report.date,
        employee_id=report.support_user_id,
        employee=UserOut(
            id=employee.id,
            username=employee.username,
            full_name=employee.full_name,
            role=employee.role,
            is_active_for_duties=bool(employee.is_active_for_duties),
            is_eligible_for_morning_duties=bool(employee.is_eligible_for_morning_duties),
            is_bootstrap_admin=False,
            bitrix_user_id=None,
        ),
        status=report.status,
        finalized_at=report.finalized_at,
        updated_at=report.updated_at,
        entries=[
            ReportEntryOut(task=e.task or "", minutes=e.minutes, description=e.description) for e in report.entries
        ],
    )


def build_qrcode_png_bytes(url: str) -> bytes:
    try:
        import qrcode  # type: ignore[import-untyped]
        import qrcode.constants  # type: ignore[import-untyped]
    except ImportError as e:  # pragma: no cover
        raise HTTPException(
            status_code=503,
            detail="Генерация QR недоступна: установите пакет qrcode[pil] и пересоберите backend.",
        ) from e
    qr = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=6,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()
